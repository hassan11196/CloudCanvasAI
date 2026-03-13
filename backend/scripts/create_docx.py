#!/usr/bin/env python3
"""
DOCX Document Generator
Usage: python create_docx.py <output_path> <json_content>

The JSON content should have this structure:
{
    "title": "Document Title",
    "sections": [
        {
            "heading": "Section Heading",
            "level": 1,  // 1 for Heading 1, 2 for Heading 2, etc.
            "content": "Paragraph text or list of paragraphs",
            "bullet_points": ["item 1", "item 2"]  // Optional
        }
    ]
}
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def create_document(output_path: str, content: dict) -> str:
    """Create a Word document from structured content."""
    doc = Document()

    # Set document title if provided
    if "title" in content:
        title = doc.add_heading(content["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process sections
    for section in content.get("sections", []):
        # Add heading
        if "heading" in section:
            level = section.get("level", 1)
            doc.add_heading(section["heading"], level=min(level, 9))

        # Add paragraph content
        if "content" in section:
            content_text = section["content"]
            if isinstance(content_text, list):
                for para in content_text:
                    doc.add_paragraph(para)
            else:
                doc.add_paragraph(content_text)

        # Add bullet points
        if "bullet_points" in section:
            for item in section["bullet_points"]:
                doc.add_paragraph(item, style="List Bullet")

        # Add numbered list
        if "numbered_list" in section:
            for item in section["numbered_list"]:
                doc.add_paragraph(item, style="List Number")

        # Add table if provided
        if "table" in section:
            table_data = section["table"]
            if table_data.get("headers") and table_data.get("rows"):
                headers = table_data["headers"]
                rows = table_data["rows"]

                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"

                # Add headers
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header

                # Add rows
                for row_data in rows:
                    row_cells = table.add_row().cells
                    for i, cell_text in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].text = str(cell_text)

    # Ensure the output directory exists
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    # Save the document
    doc.save(output_path)
    return output_path


def validate_content(content: dict) -> None:
    """Validate document content structure.

    Args:
        content: Document content dictionary

    Raises:
        ValueError: If content structure is invalid
    """
    if not isinstance(content, dict):
        raise ValueError("Content must be a dictionary")

    # Validate sections if present
    sections = content.get("sections", [])
    if not isinstance(sections, list):
        raise ValueError("Sections must be a list")

    if len(sections) > 1000:
        raise ValueError("Too many sections (max 1000)")

    for idx, section in enumerate(sections):
        if not isinstance(section, dict):
            raise ValueError(f"Section {idx} must be a dictionary")

        # Validate section content size
        if "content" in section:
            content_text = section["content"]
            if isinstance(content_text, str) and len(content_text) > 100000:
                raise ValueError(f"Section {idx} content too large (max 100KB)")
            elif isinstance(content_text, list) and len(content_text) > 100:
                raise ValueError(f"Section {idx} has too many paragraphs (max 100)")


def main():
    if len(sys.argv) < 3:
        print("Usage: python create_docx.py <output_path> <json_content>")
        print("Or: python create_docx.py <output_path> --file <json_file>")
        sys.exit(1)

    try:
        output_path = sys.argv[1]

        # Check if reading from file
        if len(sys.argv) >= 4 and sys.argv[2] == "--file":
            json_file = sys.argv[3]
            # Check file size before reading
            file_size = Path(json_file).stat().st_size
            if file_size > 10 * 1024 * 1024:  # 10 MB
                print(f"Error: JSON file too large ({file_size} bytes, max 10MB)")
                sys.exit(1)

            with open(json_file, "r", encoding="utf-8") as f:
                content = json.load(f)
        else:
            # Parse JSON from command line
            json_str = sys.argv[2]
            # Limit JSON string size
            if len(json_str) > 1024 * 1024:  # 1 MB
                print(f"Error: JSON string too large ({len(json_str)} bytes, max 1MB)")
                sys.exit(1)

            content = json.loads(json_str)

        # Validate content structure
        validate_content(content)

        result = create_document(output_path, content)
        print(f"Document created: {result}")

    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON - {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"Error: Invalid content - {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
